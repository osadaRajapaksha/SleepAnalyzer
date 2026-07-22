import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_paragraph(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    return p

def add_claim(doc, number, text):
    p = doc.add_paragraph(f"{number}. {text}")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    return p

def create_patent_doc(filename):
    doc = Document()
    
    # Description Main Heading
    add_heading(doc, 'DESCRIPTION', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title of the invention
    add_heading(doc, 'Title of the invention', level=2)
    add_paragraph(doc, 'MULTI-TARGET PHYSIOLOGICAL MONITORING SYSTEM')

    # Technical Field
    add_heading(doc, 'Technical Field', level=2)
    add_paragraph(doc, 
        "The present invention relates generally to the field of non-invasive physiological monitoring. More specifically, "
        "the invention relates to signal processing architectures for multi-occupant environments, providing systems and methods "
        "for recovering, attributing, and maintaining the continuous identity of latent physiological signals extracted from "
        "mixed sensor arrays.", indent=True)

    # Background art
    add_heading(doc, 'Background art', level=2)
    add_paragraph(doc, 
        "Non-invasive physiological monitoring, particularly sleep monitoring using under-mattress sensors, has gained significant "
        "popularity. Such systems utilize force sensors, piezo-electric films, or strain gauges to detect ballistocardiographic "
        "(BCG) and respiratory efforts. However, when monitoring multiple subjects in close proximity (e.g., a double bed), "
        "the sensors inexorably capture a mixed signal containing overlapping physiological data and noise from both subjects.", indent=True)

    # Technical problem
    add_heading(doc, 'Technical problem', level=2)
    add_paragraph(doc,
        "Prior art systems attempt to resolve this cross-talk by relying on discrete movement events. For example, traditional "
        "methods determine which subject moved by comparing the signal amplitude of a sudden movement and measuring the time delay "
        "between the movement registering on a first sensor versus a second sensor. If the movement registers stronger and earlier "
        "on the left sensor, it is attributed to the left occupant. While functional for isolated movements, these methodologies "
        "fail dramatically during continuous periods of rest, simultaneous movement, or when occupants roll over into each other's "
        "respective sensor zones. Because traditional systems do not disentangle the continuous underlying physiological waveforms "
        "(i.e., the latent sources), they are highly susceptible to 'identity swapping'—where the system mistakenly assigns the "
        "heart rate of Subject A to Subject B following a complex physical shift. Therefore, there exists a profound need in the art for a physiological monitoring system capable of blindly separating "
        "mixed physiological signals into continuous latent source streams, assigning those streams with probabilistic confidence, "
        "and persistently tracking identities to ensure high-fidelity longitudinal data.", indent=True)

    # Technical Solution
    add_heading(doc, 'Technical Solution', level=2)
    add_paragraph(doc, 
        "The present invention provides a robust, three-engine signal processing architecture "
        "capable of isolating and continuously tracking physiological signatures from multiple occupants in a shared environment. "
        "A Source Recovery Engine receives raw, mixed signal streams from a multi-channel "
        "sensor array and applies a Blind Source Separation (BSS) algorithm, such as Independent Component Analysis (ICA), across sliding time windows "
        "to mathematically reconstruct the latent, independent physiological waveforms of each occupant. "
        "An Occupant Attribution Engine analyzes the reconstructed latent waveforms to determine a spatial footprint, fusing this with respiration "
        "synchronization and signal variance metrics. The engine outputs an identity assignment accompanied by a calculated "
        "Attribution Confidence Score. An Identity Persistence Engine acts as a temporal state machine that enforces a hysteresis constraint wherein "
        "a proposed identity swap is rejected unless the Attribution "
        "Confidence Score strictly exceeds a predefined high-confidence threshold. This effectively eliminates erroneous identity "
        "swapping caused by transient noise or simultaneous movement.", indent=True)

    # Brief description of drawings
    add_heading(doc, 'Brief description of drawings', level=2)
    add_paragraph(doc, "FIG. 1 illustrates an overview of the multi-target physiological monitoring system 100.", indent=True)
    add_paragraph(doc, "FIG. 2 shows a block diagram of the three-engine signal processing architecture.", indent=True)
    add_paragraph(doc, "FIG. 3 illustrates the operations within the Data Acquisition Layer 110.", indent=True)
    add_paragraph(doc, "FIG. 4 is a flow diagram detailing the Source Recovery Engine 200 operations.", indent=True)
    add_paragraph(doc, "FIG. 5 illustrates the application of the Blind Source Separation algorithm 204.", indent=True)
    add_paragraph(doc, "FIG. 6 is a block diagram of the Occupant Attribution Engine 300 logic.", indent=True)
    add_paragraph(doc, "FIG. 7 illustrates the Identity Persistence Engine 400 state machine.", indent=True)
    add_paragraph(doc, "FIG. 8 is a flowchart of a method for continuous physiological monitoring.", indent=True)

    # Advantageous effects
    add_heading(doc, 'Advantageous effects', level=2)
    add_paragraph(doc, 
        "The present invention provides highly accurate continuous physiological monitoring for multiple subjects without requiring wearable sensors. "
        "By recovering latent sources mathematically, it eliminates the false identity swaps common in prior art systems that rely merely on movement heuristics. "
        "The rigorous state-machine persistence ensures that longitudinal metrics such as sleep stages and resting heart rates are accurately attributed "
        "to the correct individual even during periods of heavy movement or close physical proximity.", indent=True)

    # Mode for invention
    add_heading(doc, 'Mode for invention', level=2)
    
    add_paragraph(doc,
        "The following detailed description illustrates embodiments of the present disclosure. These embodiments are described "
        "in sufficient detail to enable those skilled in the art to practice the invention, and it is to be understood that other "
        "embodiments may be utilized and that logical, architectural, and algorithmic changes may be made without departing from "
        "the spirit or scope of the present invention. Reference is made to FIGS. 1-8 which illustrate various aspects of the invention.", indent=True)

    add_paragraph(doc,
        "Referring to FIG. 1 and FIG. 2, the multi-target physiological monitoring system 100 comprises a sensor array 101 configured for placement in proximity to the subjects, for example PERSON A and PERSON B. In a preferred embodiment, "
        "the sensor array 101 comprises at least a first sensor 201 (S1), a second sensor 202 (S2), and optionally a third sensor 203 (S3) "
        "disposed beneath a mattress. The signals from the sensor array 101 undergo analog signal conditioning 102. An analog-to-digital converter (ADC) 103 samples the signals at a high frequency. "
        "The raw data is continuously buffered and passed to a processing unit 104.", indent=True)

    add_paragraph(doc,
        "As shown in FIGS. 3 and 4, the buffered mixed signals from the sensors, such as mixed signals 401, 402, and 403, are passed to a signal separation engine 404 (corresponding to the signal separation module 105). The signal separation engine 404 executes "
        "a Blind Source Separation (BSS) algorithm. As illustrated by equations 306, the sensors 303, 304, 305 record a linear mixture of the latent sources, such as heartbeat A 301 and heartbeat B 302. The signal separation engine 404 mathematically decomposes "
        "the mixed multi-channel signals into a set of maximally independent latent source components. In a preferred embodiment, "
        "this generates a set of continuous, disentangled waveforms corresponding "
        "to the individual physiological rhythms, such as a recovered heartbeat A 405 corresponding to PERSON A, and a recovered heartbeat B 406 corresponding to PERSON B.", indent=True)

    add_paragraph(doc,
        "Referring back to FIG. 1, because BSS algorithms inherently suffer from permutation ambiguity, an occupant identification module 108 must resolve which latent source corresponds to which physical occupant. The occupant identification module 108 calculates a spatial "
        "fingerprint for the recovered heartbeats 405, 406 by determining cross-correlations with the raw, unfiltered signals originating from specific "
        "physical locations (e.g., S1 201 vs. S3 203). By fusing spatial correlation and signal energy, the occupant identification module 108 formulates a probabilistic assignment matrix. "
        "A heartbeat detection module 106 and heart rate calculation module 107 then extract the final physiological metrics. The metrics are transmitted via a wireless communication module 109 to a user interface 110.", indent=True)

    add_paragraph(doc,
        "Continuous physiological monitoring is highly susceptible to momentary artifacts. To prevent "
        "transient assignment errors from corrupting longitudinal sleep metrics, the processing unit 104 maintains a running memory of the active "
        "identity mapping. If a mapping is proposed that contradicts the active identity mapping (e.g., suggesting "
        "PERSON A is now on the sensor 203), the system evaluates an attribution confidence score against a "
        "pre-configured swap-threshold. If the threshold is not met, the proposed swap is rejected as noise, "
        "and the active identity mapping is persisted. If the threshold is exceeded, the swap is approved, accommodating legitimate physical "
        "crossover events without losing the continuous physiological thread of either subject.", indent=True)

    # Industrial Applicability
    add_heading(doc, 'Industrial Applicability', level=2)
    add_paragraph(doc, 
        "The invention is industrially applicable in the manufacturing of medical devices, consumer health monitoring products, and smart home appliances. "
        "It can be implemented as a commercial under-mattress sleep monitor, integrated directly into smart beds, or utilized in clinical settings for non-obtrusive patient monitoring.", indent=True)

    doc.add_page_break()

    # Claims
    add_heading(doc, 'Claims', level=1)
    add_paragraph(doc, "What is claimed is:")
    
    claims = [
        "A multi-target physiological monitoring system, comprising:\n"
        "a sensor array configured to acquire a plurality of mixed physiological signals;\n"
        "a processor communicatively coupled to the sensor array;\n"
        "a non-transitory computer-readable medium storing instructions that, when executed by the processor, cause the system to:\n"
        "  buffer the acquired mixed physiological signals into a temporal window;\n"
        "  isolate a plurality of independent latent physiological waveforms from the temporal window using blind source separation;\n"
        "  generate a spatial footprint for each isolated latent physiological waveform to determine an identity assignment mapping and a confidence metric; and\n"
        "  persist a historical identity assignment mapping across sequential temporal windows, wherein the system rejects a change to the historical identity assignment mapping unless the confidence metric surpasses a defined swapping threshold.",
        
        "The multi-target physiological monitoring system of claim 1, wherein isolating the plurality of independent latent physiological waveforms is performed exclusively on continuous resting physiological data without requiring the detection of a movement event amplitude or time delay.",
        
        "The multi-target physiological monitoring system of claim 1, wherein the sensor array is adapted to be positioned beneath a mattress, and the independent latent physiological waveforms correspond to individual ballistocardiogram signatures of occupants sharing the mattress.",

        "A method for continuous, multi-target physiological monitoring in a shared environment, the method comprising:\n"
        "receiving, from a sensor array, a plurality of mixed physiological signals spanning a temporal window;\n"
        "reconstructing, via a processor executing a Source Recovery Engine, a plurality of continuous latent physiological source streams from the mixed physiological signals, wherein the reconstruction is performed independently of discrete movement event detection;\n"
        "calculating, via an Occupant Attribution Engine, an attribution mapping and a corresponding confidence score for each reconstructed latent physiological source stream based on a correlation between the reconstructed stream and the plurality of mixed physiological signals; and\n"
        "maintaining, via an Identity Persistence Engine, a longitudinal identity association for each target by enforcing a state machine constraint, wherein an existing longitudinal identity association is overridden only if the calculated confidence score exceeds a predefined threshold.",
        
        "The method of claim 4, wherein reconstructing the plurality of continuous latent physiological source streams comprises applying a Blind Source Separation (BSS) algorithm to the mixed physiological signals.",
        
        "The method of claim 5, wherein the Blind Source Separation (BSS) algorithm is Independent Component Analysis (ICA).",
        
        "The method of claim 4, further comprising applying a bandpass filter to the mixed physiological signals prior to reconstructing the latent physiological source streams, wherein the bandpass filter isolates frequencies corresponding to human cardiac or respiratory activity.",
        
        "The method of claim 4, wherein calculating the attribution mapping comprises generating a spatial footprint by calculating a cross-correlation matrix between the continuous latent physiological source streams and the raw mixed physiological signals.",
        
        "The method of claim 4, wherein maintaining the longitudinal identity association comprises rejecting an attribution mapping that swaps target identities if the confidence score falls below 80 percent.",
        
        "The method of claim 4, wherein the sensor array comprises at least two unobtrusive sensors selected from the group consisting of strain gauges, piezoelectric elements, and hydraulic sensors."
    ]

    for i, claim_text in enumerate(claims, 1):
        add_claim(doc, i, claim_text)

    doc.add_page_break()
    
    # Abstract
    add_heading(doc, 'Abstract', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 
        "A system and method for continuous, multi-target physiological monitoring that separates and persistently tracks "
        "signals from multiple subjects sharing a monitored environment. The system comprises a sensor array to capture a "
        "plurality of mixed physiological signals. A Source Recovery Engine actively reconstructs continuous latent physiological "
        "streams from the mixed signals using blind source separation techniques, bypassing reliance on discrete movement events. "
        "An Occupant Attribution Engine assigns the recovered streams to respective subjects based on a multi-dimensional spatial "
        "and temporal confidence score. An Identity Persistence Engine maintains longitudinal identity associations using a state "
        "machine that overrides existing identity mappings only if an attribution confidence metric surpasses a predefined high-confidence "
        "threshold. The invention prevents identity swapping during simultaneous movement or sensor zone crossover.", 
        indent=True)

    # Drawings Section
    image_files = [
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM1.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM2.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM3.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM4.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM5.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM6.png',
        'ChatGPT Image Jul 22, 2026, 06_18_41 AM7.png'
    ]
    
    for i, img_file in enumerate(image_files, 1):
        if i % 2 != 0:
            doc.add_page_break()
        if os.path.exists(img_file):
            doc.add_picture(img_file, height=Inches(3.5))
            p = doc.add_paragraph(f'FIG. {i}')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            if i % 2 != 0:
                p.paragraph_format.space_after = Pt(24)
        else:
            print(f"Warning: Could not find {img_file}")

    doc.save(filename)
    print(f"Successfully generated better patent doc at {filename}")

if __name__ == '__main__':
    create_patent_doc(r"d:\SleepAnalyzer\Professional_Patent_Draft_v2.docx")
